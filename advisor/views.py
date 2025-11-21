import google.generativeai as genai
from django.conf import settings
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login as auth_login
from django.shortcuts import HttpResponseRedirect
from django.contrib.auth.models import User
from .forms import SimpleSignupForm
from django.utils.html import escape
from django.utils.safestring import mark_safe

from .models import Conversation, Message, UploadedFile
from .forms import UploadCVForm

genai.configure(api_key=getattr(settings, "GOOGLE_API_KEY", None))


def format_ai_response(text):
    """Format model plain text into simple HTML (headings, lists, paragraphs, ordered lists)."""
    from django.utils.html import escape
    from django.utils.safestring import mark_safe
    if not text:
        return ""
    lines = text.splitlines()
    out = []
    in_list = False
    list_type = None
    for line in lines:
        s = line.strip()
        if not s:
            if in_list:
                out.append('</ol>' if list_type == 'ol' else '</ul>')
                in_list = False
                list_type = None
            continue
        if s.endswith(":") and s.split()[0].rstrip(':') in ("Summary", "Subtopics", "Details", "Next", "Next Steps", "Roadmap", "ATS", "ATS Assessment", "Top Job Suggestions", "Job Readiness"):
            if in_list:
                out.append('</ol>' if list_type == 'ol' else '</ul>')
                in_list = False
                list_type = None
            out.append(f"<h4 class=\"ai-heading\">{escape(s.rstrip(':'))}</h4>")
            continue
        # numbered roadmap lines like '1. Step description...'
        if s and s[0].isdigit() and s.split('.', 1)[0].isdigit() and s.split('.', 1)[1].strip().startswith(' '):
            if not in_list:
                out.append('<ol class="ai-roadmap">')
                in_list = True
                list_type = 'ol'
            content = s.split('.', 1)[1].strip()
            out.append(f"<li>{escape(content)}</li>")
            continue
        if s.startswith("- ") or s.startswith("* "):
            if not in_list:
                out.append('<ul class="ai-bullets">')
                in_list = True
                list_type = 'ul'
            out.append(f"<li>{escape(s[2:])}</li>")
            continue
        out.append(f"<p>{escape(s)}</p>")
    if in_list:
        out.append('</ol>' if list_type == 'ol' else '</ul>')
    return mark_safe('\n'.join(out))


def extract_text_from_file(path):
    """Extract text from txt, pdf, and docx files. Returns None on failure."""
    import os
    name = os.path.basename(path).lower()
    if name.endswith('.txt'):
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as fh:
                return fh.read()
        except Exception:
            return None
    if name.endswith('.pdf'):
        try:
            import PyPDF2
        except Exception:
            return None
        try:
            text_parts = []
            with open(path, 'rb') as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    try:
                        text_parts.append(page.extract_text() or '')
                    except Exception:
                        pass
            return '\n'.join(text_parts)
        except Exception:
            return None
    if name.endswith('.docx') or name.endswith('.doc'):
        try:
            import docx
        except Exception:
            return None
        try:
            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            return '\n'.join(paragraphs)
        except Exception:
            return None
    return None


def _get_or_create_conversation(request):
    # Switch conversation if requested
    conv_id = request.GET.get("conversation_id")
    if conv_id:
        try:
            conv = Conversation.objects.get(id=int(conv_id))
            request.session["conversation_id"] = conv.id
            return conv
        except Conversation.DoesNotExist:
            pass

    # Clear all history if requested
    if request.GET.get("clear"):
        Conversation.objects.all().delete()

    # New conversation requested
    if request.GET.get("new"):
        conv = Conversation.objects.create()
        # Immediately add greeting AI message and set state
        greeting = "Hello — please enter which career you are interested in (brief)."
        conv.state = "await_interest"
        conv.save()
        Message.objects.create(conversation=conv, sender="ai", text=greeting)
        request.session["conversation_id"] = conv.id
        return conv

    # Use session-stored conversation or create one
    conv_id = request.session.get("conversation_id")
    if conv_id:
        try:
            return Conversation.objects.get(id=conv_id)
        except Conversation.DoesNotExist:
            pass

    conv = Conversation.objects.create()
    request.session["conversation_id"] = conv.id
    return conv


@login_required
def home(request):

    response_text = ""
    conversation = _get_or_create_conversation(request)

    if request.method == "POST":
        user_input = (request.POST.get("user_input") or "").strip()
        Message.objects.create(conversation=conversation, sender="user", text=user_input)
        if not conversation.title:
            title = user_input
            if title:
                conversation.title = title[:200]
                conversation.save()

        model = genai.GenerativeModel("gemini-2.5-flash")
        response_text = ""
        try:
            def structured_prompt_for_interest(area):
                return (
                    f"You are a helpful career advisor. The user is interested in '{area}'. "
                    "Provide the response using these labeled sections exactly:\n"
                    "Summary:\nSubtopics:\nDetails:\nNext Steps:\n\n"
                    "Under 'Subtopics:' list 5 concise bullet subtopics (one per line starting with '- '). "
                    "At the end, ask the user which subtopic they'd like more details on. Reply in clear sentences. "
                    "Also include one brief follow-up question (<=20 words) at the very end to keep the conversation going."
                )

            def structured_prompt_for_subtopic(topic, area):
                return (
                    f"You are a detailed career advisor. The user previously asked about '{area}' and now selected the subtopic '{topic}'.\n"
                    "Provide ONLY the following labeled sections exactly (use these labels and formatting):\n"
                    "Details:\nRoadmap:\n\n"
                    "In 'Details:' give a clear explanation of the subtopic: what it is and relevant career paths. "
                    "In 'Roadmap:' provide a numbered list of at least 5 practical steps the user can follow to pursue careers in this subtopic. "
                    "Format each roadmap step as a numbered line starting with '1. ' (e.g. '1. Learn the basics (2-4 weeks): ...'). "
                    "Include an estimated duration for each step in parentheses (e.g. '2-4 weeks', '3-6 months'). "
                    "Where appropriate include one recommended resource or action per step. Reply in clear sentences and do not repeat previous subtopics or other sections. "
                    "Include one brief follow-up question (<=20 words) at the very end to prompt the user for the next input."
                )

            if conversation.state in ("new", "await_interest"):
                prompt = (
                    structured_prompt_for_interest(user_input)
                    + "\nPlease provide more details about which subtopic you'd like me to expand on."
                )
                response = model.generate_content(prompt)
                response_text = response.text
                conversation.state = "asked_subtopics"
                conversation.save()
            elif conversation.state == "asked_subtopics":
                prompt = structured_prompt_for_subtopic(user_input, conversation.title)
                response = model.generate_content(prompt)
                response_text = response.text
                conversation.state = "detailed"
                conversation.save()
            else:
                prompt = (
                    f"You are a concise career advisor. Provide a short structured response for: {user_input}. "
                    "Use the labeled sections: Summary:, Subtopics:, Details:, Next Steps:. "
                    "At the end include one short follow-up question (<=20 words) to keep the user engaged."
                )
                response = model.generate_content(prompt)
                response_text = response.text
        except Exception as e:
            response_text = f"Error: {e}"

        def format_ai_response(text):
            if not text:
                return ""
            lines = text.splitlines()
            out = []
            in_list = False
            for line in lines:
                s = line.strip()
                if not s:
                    if in_list:
                        out.append("</ul>")
                        in_list = False
                    continue
                if s.endswith(":") and s.split()[0].rstrip(':') in ("Summary", "Subtopics", "Details", "Next", "Next Steps", "Roadmap"):
                    if in_list:
                        out.append("</ul>")
                        in_list = False
                    out.append(f"<h4 class=\"ai-heading\">{escape(s.rstrip(':'))}</h4>")
                    continue
                if s and s[0].isdigit() and s.split('.', 1)[0].isdigit() and s.split('.', 1)[1].strip().startswith(' '):
                    if not in_list:
                        out.append('<ol class="ai-roadmap">')
                        in_list = True
                    content = s.split('.', 1)[1].strip()
                    out.append(f"<li>{escape(content)}</li>")
                    continue
                if s.startswith("- ") or s.startswith("* "):
                    if not in_list:
                        out.append("<ul class=\"ai-bullets\">")
                        in_list = True
                    out.append(f"<li>{escape(s[2:])}</li>")
                    continue
                out.append(f"<p>{escape(s)}</p>")
            if in_list:
                last = out[-1] if out else ''
                if last.startswith('<li>') or '<ol' in '\n'.join(out):
                    out.append('</ol>' if '<ol' in '\n'.join(out) else '</ul>')
            return mark_safe('\n'.join(out))

        formatted = format_ai_response(response_text)
        Message.objects.create(conversation=conversation, sender="ai", text=formatted)

        # AJAX support: return JSON if request is AJAX
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            from django.http import JsonResponse
            return JsonResponse({"reply": response_text})
        # Otherwise, redirect to GET
        return redirect("home")

    # show only conversations that have at least one message (actual history)
    conversations = Conversation.objects.filter(messages__isnull=False).distinct().order_by("-created_at")[:20]
    messages = conversation.messages.all()

    return render(request, "advisor/home.html", {"conversation": conversation, "messages": messages, "conversations": conversations})


@login_required
def upload_cv(request):
    """Simple CV upload + AI analysis endpoint."""
    result_html = None
    form = UploadCVForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        f = form.cleaned_data['file']
        # save uploaded file record
        uploaded = UploadedFile.objects.create(file=f, uploaded_by=request.user)
        path = uploaded.file.path

        # attempt extraction with debug information to give better feedback
        def try_extract_debug(path):
            import os
            name = os.path.basename(path).lower()
            reasons = []
            # txt
            if name.endswith('.txt'):
                try:
                    with open(path, 'r', encoding='utf-8', errors='ignore') as fh:
                        return fh.read(), None
                except Exception as e:
                    reasons.append(f"TXT read error: {e}")
            # pdf
            if name.endswith('.pdf'):
                try:
                    import PyPDF2
                except Exception as e:
                    reasons.append(f"PyPDF2 not available: {e}")
                else:
                    try:
                        text_parts = []
                        with open(path, 'rb') as fh:
                            reader = PyPDF2.PdfReader(fh)
                            for page in reader.pages:
                                try:
                                    text_parts.append(page.extract_text() or '')
                                except Exception:
                                    pass
                        joined = '\n'.join(text_parts)
                        # detect if extraction returned PDF internal structure (common when page.extract_text fails)
                        sample = joined[:1000]
                        def looks_like_pdf_internal(s):
                            if not s:
                                return True
                            markers = ['%PDF', '/Type', '/Font', 'xref', 'stream']
                            for m in markers:
                                if m in s:
                                    return True
                            # printable ratio
                            printable = sum(1 for c in s if 32 <= ord(c) <= 126)
                            if len(s) > 0 and (printable / max(1, len(s))) < 0.6:
                                return True
                            return False

                        if joined and not looks_like_pdf_internal(sample):
                            return joined, None
                        # fallback: try pdfminer.six if PyPDF2 result looks like internal PDF data or empty
                        try:
                            from pdfminer.high_level import extract_text as pdfminer_extract_text
                        except Exception as e:
                            reasons.append(f"PyPDF2 output unusable and pdfminer.six not available: {e}")
                        else:
                            try:
                                pm_text = pdfminer_extract_text(path)
                                if pm_text and not looks_like_pdf_internal(pm_text[:1000]):
                                    return pm_text, None
                                else:
                                    reasons.append("pdfminer.six returned no usable text")
                            except Exception as e:
                                reasons.append(f"pdfminer extraction error: {e}")
                    except Exception as e:
                        reasons.append(f"PDF extraction error: {e}")
            # docx
            if name.endswith('.docx') or name.endswith('.doc'):
                try:
                    import docx
                except Exception as e:
                    reasons.append(f"python-docx not available: {e}")
                else:
                    try:
                        doc = docx.Document(path)
                        paragraphs = [p.text for p in doc.paragraphs]
                        return '\n'.join(paragraphs), None
                    except Exception as e:
                        reasons.append(f"DOCX extraction error: {e}")

            # fallback: try reading as binary and attempt utf-8 decode
            try:
                with open(path, 'rb') as fh:
                    raw = fh.read()
                    try:
                        return raw.decode('utf-8', errors='ignore'), None
                    except Exception as e:
                        reasons.append(f"Binary decode error: {e}")
            except Exception as e:
                reasons.append(f"Could not open file: {e}")

            return None, '<br>'.join(reasons) if reasons else 'Unknown error'

        text, error = try_extract_debug(path)
        if not text:
            msg = "Could not extract text from this file."
            if error:
                msg += f" Details: {error}"
            result_html = f"<div style='color:#b91c1c;background:#fff1f2;padding:12px;border-radius:8px'>{escape(msg)}</div>"
            return render(request, 'advisor/upload.html', {'form': form, 'result_html': result_html})

        # build structured prompt
        intro = text if len(text) <= 12000 else text[:12000]
        prompt = (
            "You are a helpful and practical career advisor. Analyze the candidate's CV below and provide ONLY the "
            "following labeled sections exactly (use these labels and formatting):\n\n"
            "ATS Assessment:\n"
            "Top Job Suggestions:\n"
            "Job Readiness:\n"
            "Roadmap:\n\n"
            "Instructions:\n"
            "- Under 'ATS Assessment:' state whether the CV is ATS-friendly (answer 'Yes' or 'No') and if 'No' give specific, actionable changes the candidate should make to make it ATS-friendly (file type, section titles, keywords, formatting, fonts, removing headers/footers, etc.).\n"
            "- Under 'Top Job Suggestions:' list up to 3 job roles or domains the candidate is best suited for; for each, give a one-line justification tied to the CV evidence. Use bullet lines starting with '- '.\n"
            "- Under 'Job Readiness:' state whether the candidate appears job-ready for the suggested roles (answer 'Yes' or 'No'). If 'No', briefly state the main gaps.\n"
            "- Under 'Roadmap:' if the candidate is not job-ready, provide a numbered roadmap of at least 5 practical steps with estimated durations (e.g., '2-4 weeks') and one recommended resource or action per step. Use numbered lines starting with '1. '.\n"
            "Keep the response concise and factual. Do NOT include anything besides the labeled sections above.\n\n"
            f"CV TEXT:\n{intro}"
        )

        model = genai.GenerativeModel("gemini-2.5-flash")
        try:
            resp = model.generate_content(prompt)
            ai_text = resp.text
        except Exception as e:
            ai_text = f"Error: {e}"

        # reuse existing formatter to turn model text into safe HTML
        result_html = format_ai_response(ai_text)

    return render(request, 'advisor/upload.html', {'form': form, 'result_html': result_html})


def signup(request):
    if request.method == 'POST':
        form = SimpleSignupForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']
            user = User.objects.create_user(username=username, password=password)
            auth_login(request, user)
            return redirect('home')
    else:
        form = SimpleSignupForm()
    return render(request, 'registration/signup.html', {'form': form})


@login_required
def stream_chat(request):
    """Stream chat responses back to the browser using SSE-style events.

    Fallback behaviour: if the model client doesn't support streaming, we generate
    the full reply and then stream it in small chunks so the frontend receives
    content progressively.
    """
    from django.http import StreamingHttpResponse

    if request.method != 'POST':
        return StreamingHttpResponse(status=405)

    user_input = (request.POST.get('user_input') or '').strip()
    if not user_input:
        return StreamingHttpResponse(status=400)

    # Save user message to DB
    conversation = _get_or_create_conversation(request)
    Message.objects.create(conversation=conversation, sender='user', text=user_input)

    model = genai.GenerativeModel("gemini-2.5-flash")

    # Build prompt similar to home view (reuse earlier helpers if necessary)
    prompt = (
        f"You are a concise career advisor. Provide a short structured response for: {user_input}. "
        "Use the labeled sections: Summary:, Subtopics:, Details:, Next Steps:. At the end include one short follow-up question (<=20 words)."
    )

    def event_stream_from_text(text, chunk_size=40, delay=0.02):
        """Yield server-sent-event formatted chunks from a full text response."""
        import time
        i = 0
        L = len(text)
        while i < L:
            chunk = text[i:i+chunk_size]
            i += chunk_size
            # SSE format: data: <payload>\n\n
            yield f"data: {chunk}\n\n"
            if delay:
                time.sleep(delay)
        # signal end
        yield "event: done\ndata: \n\n"

    # Try to use streaming API if available. If not, fall back to generating full text
    try:
        # Some clients expose a streaming generator. Try to detect common APIs.
        if hasattr(model, 'stream'):
            # hypothetical streaming API: model.stream(prompt) yields partial strings
            def generator():
                try:
                    for part in model.stream(prompt):
                        yield f"data: {part}\n\n"
                    yield "event: done\ndata: \n\n"
                except Exception as e:
                    yield f"event: error\ndata: {escape(str(e))}\n\n"
            return StreamingHttpResponse(generator(), content_type='text/event-stream')

        # Another hypothetical API: model.generate_content(..., stream=True) returns iterable
        if hasattr(model, 'generate_content'):
            # Some SDKs support stream=True; try and catch if unsupported
            try:
                gen = model.generate_content(prompt, stream=True)
                if hasattr(gen, '__iter__'):
                    def generator_from_gen():
                        try:
                            for part in gen:
                                # part might be object with .text or a string
                                text_part = getattr(part, 'text', str(part))
                                yield f"data: {text_part}\n\n"
                            yield "event: done\ndata: \n\n"
                        except Exception as e:
                            yield f"event: error\ndata: {escape(str(e))}\n\n"
                    return StreamingHttpResponse(generator_from_gen(), content_type='text/event-stream')
            except TypeError:
                # stream not supported, fall back
                pass
            except Exception:
                # any other failure, we'll fall back
                pass

        # Fallback: generate full response and stream it in chunks
        response = model.generate_content(prompt)
        full_text = response.text

        # Save AI response as before (formatted HTML)
        def format_ai_response_for_save(text):
            # reuse simpler formatting used elsewhere
            return format_ai_response(text)

        Message.objects.create(conversation=conversation, sender='ai', text=format_ai_response_for_save(full_text))

        return StreamingHttpResponse(event_stream_from_text(full_text), content_type='text/event-stream')

    except Exception as e:
        # If something unexpected happens, return an error event
        def err():
            yield f"event: error\ndata: {escape(str(e))}\n\n"
        return StreamingHttpResponse(err(), content_type='text/event-stream')
